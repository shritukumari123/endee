from __future__ import annotations
import io, os, re
from dataclasses import dataclass
from loguru import logger
from core.models import FileType

POPPLER_PATH = r"C:\Users\Admin\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


@dataclass
class ParsedFile:
    filename: str
    file_type: FileType
    text: str
    page_count: int = 1
    size_bytes: int = 0


def detect_file_type(filename: str) -> FileType:
    ext = os.path.splitext(filename)[1].lower()
    return {
        ".pdf":  FileType.PDF,
        ".txt":  FileType.TEXT,
        ".md":   FileType.MARKDOWN,
        ".csv":  FileType.CSV,
        ".docx": FileType.DOCX,
    }.get(ext, FileType.UNKNOWN)


def _ocr_pdf(data: bytes) -> str:
    """
    OCR fallback for scanned PDFs.
    Converts each page to image then extracts text using Tesseract.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes

        # Set Tesseract path
        if os.path.exists(TESSERACT_PATH):
            pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

        # Convert PDF pages to images using poppler
        images = convert_from_bytes(
            data,
            dpi=200,
            poppler_path=POPPLER_PATH,
        )

        texts = []
        for i, img in enumerate(images):
            logger.info(f"OCR processing page {i + 1}/{len(images)}...")
            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                texts.append(text.strip())

        full_text = "\n\n".join(texts)
        logger.success(f"OCR complete — extracted {len(full_text):,} characters")
        return full_text

    except ImportError as exc:
        logger.error(f"OCR packages missing: {exc}. Run: pip install pytesseract pdf2image pillow")
        return ""
    except Exception as exc:
        logger.error(f"OCR failed: {exc}")
        return ""


def _parse_pdf(data: bytes) -> tuple:
    """
    Parse PDF — tries normal text extraction first.
    If PDF is scanned/image-based switches to OCR automatically.
    """
    try:
        import fitz
        doc   = fitz.open(stream=data, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()

        full_text = "\n\n".join(pages).strip()

        # If less than 50 characters extracted — it is a scanned PDF
        if not full_text or len(full_text) < 50:
            logger.info("Scanned PDF detected — switching to OCR mode")
            full_text = _ocr_pdf(data)

        return full_text, len(pages)

    except ImportError:
        raise ImportError("Run: pip install PyMuPDF")
    except Exception as exc:
        logger.error(f"PDF parse error: {exc}")
        return "", 0


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text.strip() for c in row.cells if c.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts).strip()
    except ImportError:
        raise ImportError("Run: pip install python-docx")
    except Exception as exc:
        logger.error(f"DOCX parse error: {exc}")
        return ""


def _parse_csv(data: bytes) -> str:
    try:
        import pandas as pd
        df     = pd.read_csv(io.BytesIO(data)).head(500)
        header = f"Columns: {', '.join(df.columns.tolist())}."
        lines  = []
        for _, row in df.iterrows():
            parts = [
                f"{col}: {val}"
                for col, val in row.items()
                if str(val).strip() and str(val).lower() != "nan"
            ]
            if parts:
                lines.append(". ".join(parts) + ".")
        return header + "\n\n" + "\n".join(lines)
    except ImportError:
        raise ImportError("Run: pip install pandas")
    except Exception as exc:
        logger.error(f"CSV parse error: {exc}")
        return ""


def _parse_text(data: bytes) -> str:
    try:
        text = data.decode("utf-8", errors="replace")
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`[^`]+`", "", text)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
        text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    except Exception as exc:
        logger.error(f"Text parse error: {exc}")
        return ""


def parse_file(filename: str, data: bytes) -> ParsedFile:
    """
    Parse any supported file into clean text.
    Supports: PDF (normal + scanned), DOCX, CSV, TXT, MD
    """
    file_type  = detect_file_type(filename)
    size_bytes = len(data)
    pages      = 1

    logger.info(f"Parsing {filename} ({file_type.value}, {size_bytes:,} bytes)")

    if file_type == FileType.PDF:
        text, pages = _parse_pdf(data)
    elif file_type == FileType.DOCX:
        text = _parse_docx(data)
    elif file_type == FileType.CSV:
        text = _parse_csv(data)
    elif file_type in (FileType.TEXT, FileType.MARKDOWN):
        text = _parse_text(data)
    else:
        raise ValueError(
            f"Unsupported file type: '{os.path.splitext(filename)[1]}'. "
            "Supported: .pdf .txt .md .csv .docx"
        )

    if not text.strip():
        logger.warning(f"No text extracted from '{filename}'.")

    logger.success(f"Parsed '{filename}' → {len(text):,} chars")

    return ParsedFile(
        filename=filename,
        file_type=file_type,
        text=text,
        page_count=pages,
        size_bytes=size_bytes,
    )